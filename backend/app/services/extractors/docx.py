from docx import Document
from typing import Dict, Any
from app.services.extractors.base import BaseExtractor


class DOCXExtractor(BaseExtractor):
    
    def extract(self, source: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        file_path = metadata.get('file_path', source)
        
        doc = Document(file_path)
        
        content = {
            'text': '',
            'title': '',
            'author': '',
            'has_images': False,
            'has_tables': False,
            'comments': [],
        }
        
        if doc.core_properties.title:
            content['title'] = doc.core_properties.title
        
        if doc.core_properties.author:
            content['author'] = doc.core_properties.author
        
        for paragraph in doc.paragraphs:
            content['text'] += paragraph.text + '\n'
        
        if doc.tables:
            content['has_tables'] = True
            for table in doc.tables:
                content['text'] += '\n[TABLE]\n'
                for row in table.rows:
                    row_text = '\t'.join([cell.text for cell in row.cells])
                    content['text'] += row_text + '\n'
                content['text'] += '[/TABLE]\n'
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                content['has_images'] = True
                break
        
        return content
    
    def extract_metadata(self, source: str) -> Dict[str, Any]:
        doc = Document(source)
        props = doc.core_properties
        
        return {
            'title': props.title or '',
            'author': props.author or '',
            'subject': props.subject or '',
            'keywords': props.keywords or '',
            'created': str(props.created) if props.created else '',
            'modified': str(props.modified) if props.modified else '',
            'last_modified_by': props.last_modified_by or '',
        }
